import re
import docx
from docx.shared import Inches

def indentate_lines(markdown_text):
    lines = markdown_text.split('\n')
    indentated = []

    for line in lines:
        indent_level = 0

        # Count the number of double spaces at the start of the line
        while line.startswith('  '):
            indent_level += 1
            line = line[2:]  # Remove the counted double space

        indentated.append({'content': line, 'indent': indent_level})

    return indentated


def preprocess(indentated_lines):
    # We're directly dealing with dictionaries inside the list here.
    # So, each 'line_info' is a dictionary with 'content' and 'indent'.
    
    # First, we process the lines for heading styles.
    heading_2_indices = set()
    for i, line_info in enumerate(indentated_lines):
        if line_info['content'].startswith("---") and i > 0 and indentated_lines[i - 1]['content'].strip():
            indentated_lines[i]['content'] = ""  # Clear the current '---' line
            heading_2_indices.add(i - 1)

    # Apply transformations on each line's content.
    for i, line_info in enumerate(indentated_lines):
        line = replace_pattern(line_info['content'])
        if i in heading_2_indices:
            line = "## " + line  # Add Markdown heading syntax for level 2 headings
        indentated_lines[i]['content'] = line  # Update the content after transformations

    return indentated_lines


def replace_pattern(s):
    s = re.sub(r'\[\^(.*?)\^\]', r'^\1', s)
    s = re.sub(r'\[(.*?)\]', r'^\1', s)
    s = re.sub(r'^\*\s', '- ', s)
    return s

def process_for_formatting(line):
    runs = []
    i = 0
    if line.startswith('^') and re.match(r"^\^\d+", line):
        match = re.match(r"^\^(\d+)", line)
        if match:
            runs.append((match.group(1), 'bold'))
            i += len(match.group(0))

    while i < len(line):
        if line[i] == '^':
            end_sup = line.find(' ', i)
            if end_sup != -1:
                runs.append((line[i:end_sup], 'normal'))
                i = end_sup
            else:
                runs.append((line[i:], 'normal'))
                i = len(line)
            continue
        elif line[i:i+3] == '***':
            end_bold_italic = line.find('***', i + 3)
            if end_bold_italic != -1:
                runs.append((line[i+3:end_bold_italic], 'bold_italic'))
                i = end_bold_italic + 3
            else:
                runs.append((line[i], 'normal'))
                i += 1
        elif line[i:i+2] == '**':
            end_bold = line.find('**', i + 2)
            if end_bold != -1:
                inner_runs = process_for_formatting(line[i+2:end_bold])
                for run_text, run_style in inner_runs:
                    if run_style == 'italic':
                        runs.append((run_text, 'bold_italic'))
                    else:
                        runs.append((run_text, 'bold'))
                i = end_bold + 2
            else:
                runs.append((line[i], 'normal'))
                i += 1
        elif line[i] == '*':
            end_italic = line.find('*', i + 1)
            if end_italic != -1:
                runs.append((line[i+1:end_italic], 'italic'))
                i = end_italic + 1
            else:
                runs.append((line[i], 'normal'))
                i += 1
        else:
            next_bold = line.find('**', i)
            next_italic = line.find('*', i)
            next_bold_italic = line.find('***', i)
            next_superscript = line.find('^', i)

            next_special = min([pos for pos in [next_bold, next_italic, next_bold_italic, next_superscript] if pos != -1], default=len(line))
            runs.append((line[i:next_special], 'normal'))
            i = next_special

    return runs

def process_for_superscript(runs):
    processed_runs = []

    for run_text, run_style in runs:
        i = 0
        while i < len(run_text):
            if run_text[i] == '^':
                match = re.match(r"\^(\d+)", run_text[i:])
                if match:
                    superscript_text = match.group(1)
                    if run_style == 'normal':
                        processed_runs.append((superscript_text, 'superscript'))
                    elif run_style == 'bold':
                        processed_runs.append((superscript_text, 'bold_superscript'))
                    elif run_style == 'italic':
                        processed_runs.append((superscript_text, 'italic_superscript'))
                    elif run_style == 'bold_italic':
                        processed_runs.append((superscript_text, 'bold_italic_superscript'))
                    i += len(match.group(0))
                    continue
            processed_runs.append((run_text[i], run_style))
            i += 1

    return processed_runs


def add_run(paragraph, text, style):
    run = paragraph.add_run(text)
    if style == 'bold':
        run.bold = True
    elif style == 'italic':
        run.italic = True
    elif style == 'bold_italic':
        run.bold = True
        run.italic = True
    elif style == 'superscript':
        run.font.superscript = True
    elif style == 'bold_superscript':
        run.bold = True
        run.font.superscript = True
    elif style == 'italic_superscript':
        run.italic = True
        run.font.superscript = True
    elif style == 'bold_italic_superscript':
        run.bold = True
        run.italic = True
        run.font.superscript = True

def process_headings(line, doc):
    # Define markdown heading levels
    headings = [
        ("# ", 1),
        ("## ", 2),
        ("### ", 3),
        ("#### ", 4),
        ("##### ", 5),
        ("###### ", 6)
    ]

    # Check if the current line starts with a markdown heading prefix
    for prefix, level in headings:
        if line.startswith(prefix):
            # Remove the markdown prefix from the heading text
            text = line[len(prefix):]
            # Add the heading to the document with the appropriate level
            doc.add_heading(text, level=level)
            return True  # Heading was processed

    return False  # No heading was found

def process_bullets(line_info, doc):
    # Check if the current line starts with the bullet point markdown
    line = line_info['content']
    indent_level = line_info['indent']
    if line.startswith("- "):
        # Extract the actual content, excluding the markdown bullet point ("- ")
        content = line[2:]

        # Depending on the indent level, we might want to adjust the style
        if indent_level == 0:
            style = 'List Bullet'
        elif indent_level == 1:
            style = 'List Bullet 2'  # Assuming 'List Bullet 2' is defined in your Word styles
        elif indent_level == 2:
            style = 'List Bullet 3'  # Assuming 'List Bullet 3' is defined in your Word styles
        else:
            style = 'List Bullet'  # Default to normal bullets for deeper indents or unexpected cases

        # Create a new bullet point with the specified style
        bullet = doc.add_paragraph(style=style)

        # Instead of adding text directly, process the content for formatting
        runs = process_for_formatting(content)
        runs = process_for_superscript(runs)
        for run_text, run_style in runs:
            add_run(bullet, run_text, run_style)  # Apply the styles per run

        return True  # The line was a bullet point and was processed

    return False  # The line was not a bullet point

def process_numbered_lists(line_info, doc):
    # Check if the current line starts with a markdown numbered list format
    line = line_info['content']
    indent_level = line_info['indent']

    # This regular expression matches lines that start with "1. ", "2. ", etc.
    # It handles multi-digit numbers and a single following space.
    match = re.match(r'(\d+\.)+\s', line)
    if match:
        # Extract the actual content, excluding the markdown numbered prefix (e.g., "1. ")
        content = line[match.end():]  # Strip the numbers and period, leave the content.
        
        #Debugging
        content += 'THIS GOT NUMBERED'

        # Depending on the indent level, we might want to adjust the style
        if indent_level == 0:
            style = 'List Number'
        elif indent_level == 1:
            style = 'List Number 2'  # Assuming 'List Number 2' is defined in your Word styles
        elif indent_level == 2:
            style = 'List Number 3'  # Assuming 'List Number 3' is defined in your Word styles
        else:
            style = 'List Number'  # Default to normal numbering for deeper indents or unexpected cases

        # Create a new numbered list item with the specified style
        number_item = doc.add_paragraph(style=style)

        # Instead of adding text directly, process the content for formatting
        runs = process_for_formatting(content)
        runs = process_for_superscript(runs)
        for run_text, run_style in runs:
            add_run(number_item, run_text, run_style)  # Apply the styles per run

        return True  # The line was a numbered list item and was processed

    return False  # The line was not a numbered list item



def process_lines(doc, lines):
    for line_info in lines:
        # If the line is a heading, a bullet point, or a numbered list, it's processed inside these functions
        if process_headings(line_info['content'], doc):
            continue
        elif process_bullets(line_info, doc):
            continue
        elif process_numbered_lists(line_info, doc):
            continue
        else:
            # If the line is a regular line (not a heading, bullet, or numbered list), create a new paragraph for it
            paragraph = doc.add_paragraph()

            # Adjust the left indent based on the indentation level.
            if line_info['indent'] > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * line_info['indent'])

            # Process the content for formatting (bold, italic, etc.)
            runs = process_for_formatting(line_info['content'])
            runs = process_for_superscript(runs)
            for run_text, run_style in runs:
                add_run(paragraph, run_text, run_style)




def markdown_to_docx(markdown_text, filename):
    doc = docx.Document()

    # Indentate lines before processing
    indentated_lines = indentate_lines(markdown_text)

    # Preprocess the lines. Note that we're passing the list of line dictionaries here,
    # not a single string.
    preprocessed_lines = preprocess(indentated_lines)

    # The remaining part of your function stays the same
    process_lines(doc, preprocessed_lines)

    filename = filename[:42]  # Ensure the filename is not excessively long
    file_path = f"{filename}.docx"
    doc.save(file_path)

    return file_path


